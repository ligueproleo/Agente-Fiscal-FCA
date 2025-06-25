# utils/processing.py

import pandas as pd
import zipfile
import io
import unicodedata
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import streamlit as st

# As funções de processamento de dados permanecem as mesmas
@st.cache_data 
def processar_zip(arquivo_zip):
    try:
        with zipfile.ZipFile(arquivo_zip, 'r') as z:
            nome_cabecalho = [nome for nome in z.namelist() if 'Cabecalho.csv' in nome][0]
            nome_itens = [nome for nome in z.namelist() if 'Itens.csv' in nome][0]
            with z.open(nome_cabecalho) as f:
                df_cabecalho = pd.read_csv(f, sep=',', decimal='.', encoding='utf-8')
            with z.open(nome_itens) as f:
                df_itens = pd.read_csv(f, sep=',', decimal='.', encoding='utf-8')
            
            df_cabecalho = limpar_nomes_colunas(df_cabecalho)
            df_itens = limpar_nomes_colunas(df_itens)
            df_completo = pd.merge(df_cabecalho, df_itens, on='chave_de_acesso', how='inner')
            
            for col in df_completo.columns:
                if 'data' in col:
                    df_completo[col] = pd.to_datetime(df_completo[col], errors='coerce')
            return df_completo
    except Exception as e:
        raise e

def limpar_nomes_colunas(df):
    cols_novas = []
    for col in df.columns:
        col = ''.join(c for c in unicodedata.normalize('NFD', col) if unicodedata.category(c) != 'Mn')
        col = col.lower().strip().replace(' ', '_').replace('/', '_').replace('-', '_')
        col = col.replace('(', '').replace(')', '').replace('.', '')
        cols_novas.append(col)
    df.columns = cols_novas
    return df

@st.cache_data
def criar_documento_word(report_items):
    """
    Gera um documento Word profissional e bem formatado.
    """
    document = Document()
    
    styles = document.styles
    styles['Title'].font.name = 'Calibri'
    styles['Title'].font.size = Pt(26)
    styles['Heading 1'].font.name = 'Calibri'
    styles['Heading 1'].font.size = Pt(16)
    styles['Heading 2'].font.name = 'Calibri'
    styles['Heading 2'].font.size = Pt(13)

    # --- Cabeçalho do Relatório ---
    document.add_heading('Relatório de Análise de Notas Fiscais', level=0)
    p_data = document.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data_geracao = datetime.now().strftime("%d de %B de %Y, %H:%M:%S")
    p_data.add_run(f'Relatório gerado em: {data_geracao}').italic = True
    
    # --- CORREÇÃO ---
    # A linha abaixo foi removida. Ela forçava o conteúdo a começar na segunda página.
    # document.add_page_break()

    # Adiciona um parágrafo em branco para dar um espaçamento antes do primeiro item.
    document.add_paragraph()

    # --- Renderização dos Itens do Relatório ---
    for item in report_items:
        try:
            titulo_item = item.get('title', 'Item de Relatório')
            document.add_heading(titulo_item, level=2)
            
            content = item['content']
            
            if item['type'] == 'qa' or item.get('category') == 'insight_ia':
                document.add_paragraph(f"Pergunta: {content['pergunta']}", style='Intense Quote')
                document.add_paragraph(f"Resposta: {content['resposta']}")

            elif item['type'] == 'dataframe':
                df_item = content['dados']
                if df_item.index.name is not None:
                    df_item = df_item.reset_index()
                if not df_item.empty:
                    t = document.add_table(df_item.shape[0] + 1, df_item.shape[1], style='Table Grid')
                    for j, col_name in enumerate(df_item.columns):
                        cell = t.cell(0, j)
                        cell.text = str(col_name)
                        cell.paragraphs[0].runs[0].font.bold = True
                    for i in range(df_item.shape[0]):
                        for j in range(df_item.shape[1]):
                            valor = df_item.values[i, j]
                            texto = f"{valor:,.2f}" if isinstance(valor, (int, float)) else str(valor)
                            t.cell(i + 1, j).text = texto
                else:
                    document.add_paragraph("Nenhum dado para exibir nesta análise.")

            elif item['type'] == 'chart':
                fig = content['fig']
                fig.update_layout(template='plotly_white')
                
                img_buffer = io.BytesIO()
                fig.write_image(img_buffer, format='png', width=900, height=500, scale=2)
                img_buffer.seek(0)
                
                document.add_picture(img_buffer, width=Inches(6.5))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_paragraph()

        except Exception as e:
            print(f"ERRO AO PROCESSAR ITEM PARA DOCX: {item.get('title', 'N/A')}. Detalhes: {e}")
            document.add_paragraph(f"Não foi possível renderizar o item: {item.get('title', 'N/A')}", style='Body Text')

    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer